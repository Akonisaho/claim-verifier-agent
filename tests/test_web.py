"""Unit tests for web/server.py — file-extraction and input-resolution helpers.

No network/Ollama calls: these test the pure extraction logic only.
"""
import io
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT / "web"))

import docx  # python-docx
import pytest
import server as web_app  # web/server.py
from werkzeug.datastructures import FileStorage


def _file_storage(filename: str, content: bytes) -> FileStorage:
    return FileStorage(stream=io.BytesIO(content), filename=filename)


def test_extract_text_from_txt():
    fs = _file_storage("doc.txt", b"revenue grew 15% this year")
    assert web_app.extract_text_from_file(fs) == "revenue grew 15% this year"


def test_extract_text_from_docx(tmp_path):
    docx_path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("headcount increased by 4% this year")
    document.save(docx_path)
    with open(docx_path, "rb") as f:
        fs = _file_storage("doc.docx", f.read())
    assert "headcount increased by 4%" in web_app.extract_text_from_file(fs)


def test_extract_text_rejects_unsupported_extension():
    fs = _file_storage("doc.exe", b"not a real doc")
    with pytest.raises(web_app.ExtractionError):
        web_app.extract_text_from_file(fs)


def test_extract_text_empty_docx_raises(tmp_path):
    docx_path = tmp_path / "empty.docx"
    docx.Document().save(docx_path)
    with open(docx_path, "rb") as f:
        fs = _file_storage("empty.docx", f.read())
    with pytest.raises(web_app.ExtractionError):
        web_app.extract_text_from_file(fs)


def test_landing_page_links_to_verify():
    client = web_app.flask_app.test_client()
    resp = client.get("/")
    assert resp.status_code == 200
    assert b'href="/verify"' in resp.data


def test_verify_get_renders_empty_form():
    client = web_app.flask_app.test_client()
    resp = client.get("/verify")
    assert resp.status_code == 200
    assert b'action="/verify"' in resp.data
    # No results yet on a fresh GET, so no export form either.
    assert b'action="/export"' not in resp.data


def test_export_returns_a_real_docx():
    client = web_app.flask_app.test_client()
    resp = client.post(
        "/export",
        data={
            "claim": "Revenue grew 15% this year.",
            "source": "Revenue grew 15% this year.",
            "baseline_verdict": "SUPPORTED",
            "baseline_reasoning": "The number matches.",
            "verifier_verdict": "SUPPORTED",
            "verifier_audit_flag": "false",
            "claims_json": '[{"claim": "Revenue grew 15% this year.", '
            '"verdict": "SUPPORTED", "context_audit_flag": false, '
            '"source_quote": "Revenue grew 15% this year.", '
            '"reasoning": "Matches."}]',
        },
    )
    assert resp.status_code == 200
    assert resp.mimetype == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # A real .docx is a zip archive - starts with the PK signature.
    assert resp.data[:2] == b"PK"
    document = docx.Document(io.BytesIO(resp.data))
    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Revenue grew 15% this year." in all_text
    assert "SUPPORTED" in all_text
