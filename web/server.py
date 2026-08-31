"""Phase 4: a thin Flask wrapper over the existing pipeline.

No new verification logic lives here — this only accepts a document claim
and source material (pasted text or an uploaded .txt/.pdf/.docx file),
extracts plain text from the upload, and calls the exact same functions
baseline.py and app.py already expose. Nothing in those two files is
changed to support this page.

Run from the repo root:
    python web/server.py

Named server.py (not app.py) deliberately: this file lives alongside
app.py's parent directory on sys.path, and two modules both named "app"
in the same process is a real, easy-to-hit import collision — not
just a style preference.
"""
from __future__ import annotations

import io
import json
import sys
from concurrent.futures import ThreadPoolExecutor
from datetime import UTC, datetime
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

import docx  # python-docx
import pdfplumber
from flask import Flask, render_template, request, send_file
from werkzeug.datastructures import FileStorage

import app as verifier  # the verifier pipeline (app.py at repo root)
import baseline  # the baseline pipeline (baseline.py at repo root)

flask_app = Flask(__name__)
flask_app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB upload cap

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class ExtractionError(Exception):
    pass


def extract_text_from_file(file_storage: FileStorage) -> str:
    """Extract plain text from an uploaded .txt, .pdf, or .docx file."""
    filename = (file_storage.filename or "").lower()
    suffix = Path(filename).suffix
    if suffix not in ALLOWED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported file type '{suffix}'. Allowed: .txt, .pdf, .docx"
        )

    if suffix == ".txt":
        return file_storage.read().decode("utf-8", errors="replace")

    if suffix == ".pdf":
        with pdfplumber.open(file_storage) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        text = "\n".join(pages).strip()
        if not text:
            raise ExtractionError("Could not extract any text from that PDF.")
        return text

    if suffix == ".docx":
        document = docx.Document(file_storage)
        text = "\n".join(p.text for p in document.paragraphs).strip()
        if not text:
            raise ExtractionError("Could not extract any text from that .docx file.")
        return text

    raise ExtractionError(f"Unsupported file type '{suffix}'.")


def resolve_input(field_name: str) -> str:
    """A field can come from a pasted textarea OR an uploaded file.
    The uploaded file wins if both are present."""
    uploaded = request.files.get(f"{field_name}_file")
    if uploaded and uploaded.filename:
        return extract_text_from_file(uploaded)
    pasted = (request.form.get(field_name) or "").strip()
    if pasted:
        return pasted
    raise ExtractionError(
        f"Provide either pasted text or an upload for '{field_name}'."
    )


def run_baseline(claim: str, source: str) -> dict:
    raw = baseline.call_ollama(baseline.build_prompt(claim, source))
    parsed = baseline.parse_response(raw)
    return {"verdict": parsed.verdict, "reasoning": parsed.reasoning}


def run_verifier(claim: str, source: str) -> dict:
    atomic_claims = verifier.extract_claims(claim)
    # Each claim's verify+audit call is independent (same source, no shared
    # state) - dispatch them to Ollama concurrently instead of one at a time.
    # Same calls, same parsing, just not serialized; order is preserved.
    with ThreadPoolExecutor(max_workers=len(atomic_claims)) as pool:
        claim_verdicts = list(
            pool.map(lambda c: verifier.verify_claim(c.text, source), atomic_claims)
        )
    doc_verdict, audit_flag = verifier.rollup(claim_verdicts)
    return {
        "verdict": doc_verdict,
        "context_audit_flag": audit_flag,
        "claims": [
            {
                "claim": cv.claim,
                "verdict": cv.verdict,
                "context_audit_flag": cv.context_audit_flag,
                "source_quote": cv.source_quote,
                "reasoning": cv.reasoning,
            }
            for cv in claim_verdicts
        ],
    }


def build_docx_report(
    claim: str, source: str, baseline_result: dict, verifier_result: dict
) -> io.BytesIO:
    """Builds a Word (.docx) report of one verify result. python-docx is
    already a dependency (used to read .docx uploads) - no new dependency
    needed for this."""
    document = docx.Document()
    document.add_heading("Claim Verifier — Results", level=1)
    generated = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    document.add_paragraph(f"Generated {generated}")

    document.add_heading("Input", level=2)
    document.add_paragraph("Claim:", style="Intense Quote")
    document.add_paragraph(claim)
    document.add_paragraph("Source:", style="Intense Quote")
    document.add_paragraph(source)

    document.add_heading("Baseline", level=2)
    p = document.add_paragraph()
    p.add_run(baseline_result["verdict"]).bold = True
    document.add_paragraph(baseline_result["reasoning"])

    document.add_heading("Verifier", level=2)
    p = document.add_paragraph()
    p.add_run(verifier_result["verdict"]).bold = True
    if verifier_result["context_audit_flag"]:
        p.add_run("  [context audit flagged]").italic = True

    for c in verifier_result["claims"]:
        document.add_heading(c["claim"], level=3)
        p = document.add_paragraph()
        p.add_run(c["verdict"]).bold = True
        if c["context_audit_flag"]:
            p.add_run("  [audit]").italic = True
        quote_p = document.add_paragraph()
        quote_p.add_run(f"Source quote: “{c['source_quote']}”").italic = True
        document.add_paragraph(c["reasoning"])

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


@flask_app.route("/export", methods=["POST"])
def export():
    """Downloads the just-computed results (posted back as hidden form
    fields from verify.html) as a Word document."""
    claim = request.form.get("claim", "")
    source = request.form.get("source", "")
    baseline_result = {
        "verdict": request.form.get("baseline_verdict", ""),
        "reasoning": request.form.get("baseline_reasoning", ""),
    }
    verifier_result = {
        "verdict": request.form.get("verifier_verdict", ""),
        "context_audit_flag": request.form.get("verifier_audit_flag") == "true",
        "claims": json.loads(request.form.get("claims_json", "[]")),
    }
    buf = build_docx_report(claim, source, baseline_result, verifier_result)
    return send_file(
        buf,
        as_attachment=True,
        download_name="claim-verifier-results.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@flask_app.route("/", methods=["GET"])
def landing():
    """Cover page: title, thesis, a single 'start verifying' call to action."""
    return render_template("landing.html")


@flask_app.route("/verify", methods=["GET", "POST"])
def verify():
    """The verify page: GET shows the empty form, POST processes it and
    renders the same template with results."""
    if request.method == "GET":
        return render_template("verify.html")

    try:
        claim = resolve_input("claim")
        source = resolve_input("source")
    except ExtractionError as exc:
        return render_template("verify.html", error=str(exc)), 400

    try:
        # Baseline and verifier don't depend on each other - run them
        # concurrently instead of back-to-back.
        with ThreadPoolExecutor(max_workers=2) as pool:
            baseline_future = pool.submit(run_baseline, claim, source)
            verifier_future = pool.submit(run_verifier, claim, source)
            baseline_result = baseline_future.result()
            verifier_result = verifier_future.result()
    except RuntimeError as exc:
        # Ollama unreachable or failed after retries.
        return render_template("verify.html", error=str(exc)), 502

    return render_template(
        "verify.html",
        claim=claim,
        source=source,
        baseline_result=baseline_result,
        verifier_result=verifier_result,
    )


if __name__ == "__main__":
    # threaded=True: a single real verification call can take minutes (CPU
    # inference). Without this, Flask's dev server is single-threaded and
    # a second request (even just reloading the page) queues behind it.
    flask_app.run(debug=False, host="127.0.0.1", port=5000, threaded=True)
